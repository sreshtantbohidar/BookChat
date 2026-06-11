#!/usr/bin/env python3
"""
Universal Document Loader
Supports: PDF (text + OCR), TXT, MD, DOCX, CSV, JSON, XLSX, XLS,
          HTML, XML, EPUB, RTF, images (PNG/JPG/TIFF via OCR)
"""

import os
import json
import csv
import io
from pathlib import Path


def load_pdf(path):
    """Load PDF — tries text extraction first, falls back to OCR for scanned pages."""
    import fitz
    doc = fitz.open(path)
    pages = []
    needs_ocr = False

    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if not text:
            needs_ocr = True
            text = _ocr_pdf_page(page)
        pages.append(f"\n--- Page {i+1} ---\n{text}")

    doc.close()

    if needs_ocr:
        print(f"  [OCR] Some pages required OCR for {path}")

    return "\n".join(pages)


def _ocr_pdf_page(page):
    """OCR a single PDF page using Tesseract."""
    try:
        import fitz
        import pytesseract
        from PIL import Image
        # Render page to image at 300 DPI
        mat = fitz.Matrix(300/72, 300/72)
        pix = page.get_pixmap(matrix=mat)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        return f"[OCR error: {e}]"


def load_txt(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def load_docx(path):
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def load_csv(path):
    """Load CSV — converts to readable text with row/column context."""
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace", newline="") as f:
        # Try to detect dialect
        sample = f.read(8192)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample)
            reader = csv.reader(f, dialect)
        except csv.Error:
            reader = csv.reader(f)

        headers = None
        for i, row in enumerate(reader):
            if i == 0:
                headers = row
                rows.append(f"[Headers] {', '.join(headers)}")
            else:
                if headers:
                    labeled = [f"{h}: {v}" for h, v in zip(headers, row) if v.strip()]
                    rows.append(f"[Row {i}] {', '.join(labeled)}")
                else:
                    rows.append(f"[Row {i}] {', '.join(row)}")

    return "\n".join(rows)


def load_json(path):
    """Load JSON — pretty-prints with structure preserved."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        data = json.load(f)
    return _json_to_text(data)


def _json_to_text(data, indent=0):
    """Convert JSON structure to readable text."""
    lines = []
    prefix = "  " * indent

    if isinstance(data, dict):
        for key, val in data.items():
            if isinstance(val, (dict, list)):
                lines.append(f"{prefix}{key}:")
                lines.append(_json_to_text(val, indent + 1))
            else:
                lines.append(f"{prefix}{key}: {val}")
    elif isinstance(data, list):
        for i, item in enumerate(data):
            if isinstance(item, (dict, list)):
                lines.append(f"{prefix}[Item {i+1}]")
                lines.append(_json_to_text(item, indent + 1))
            else:
                lines.append(f"{prefix}- {item}")
    else:
        lines.append(f"{prefix}{data}")

    return "\n".join(lines)


def load_xlsx(path):
    """Load Excel (.xlsx) — extracts all sheets."""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            row_str = " | ".join([str(cell) if cell is not None else "" for cell in row])
            if row_str.strip():
                rows.append(f"[Row {i+1}] {row_str}")
        sheets_text.append(f"\n=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
    wb.close()
    return "\n".join(sheets_text)


def load_xls(path):
    """Load old Excel (.xls) format."""
    import xlrd
    wb = xlrd.open_workbook(path)
    sheets_text = []
    for idx in range(wb.nsheets):
        ws = wb.sheet_by_index(idx)
        rows = []
        for i in range(ws.nrows):
            row_vals = [str(ws.cell_value(i, j)) for j in range(ws.ncols)]
            row_str = " | ".join(row_vals)
            if row_str.strip():
                rows.append(f"[Row {i+1}] {row_str}")
        sheets_text.append(f"\n=== Sheet: {ws.name} ===\n" + "\n".join(rows))
    return "\n".join(sheets_text)


def load_html(path):
    """Load HTML — extracts visible text, strips tags."""
    from bs4 import BeautifulSoup
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    # Remove script/style
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    # Extract text with structure
    lines = []
    for elem in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "span", "div", "article", "section"]):
        text = elem.get_text(strip=True)
        if text and len(text) > 2:
            if elem.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                lines.append(f"\n{'#' * int(elem.name[1])} {text}\n")
            elif elem.name == "li":
                lines.append(f"  - {text}")
            else:
                lines.append(text)

    return "\n".join(lines) if lines else soup.get_text(separator="\n", strip=True)


def load_xml(path):
    """Load XML — extracts text content preserving structure."""
    from bs4 import BeautifulSoup
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "lxml-xml")

    lines = []
    for elem in soup.find_all(True):
        text = elem.get_text(strip=True)
        if text and len(text) > 2 and not elem.find():
            lines.append(f"[{elem.name}] {text}")
    return "\n".join(lines)


def load_epub(path):
    """Load EPUB — extracts text from all chapters."""
    import zipfile
    from bs4 import BeautifulSoup

    chapters = []
    with zipfile.ZipFile(path) as z:
        # Get reading order from OPF
        opf_name = None
        for name in z.namelist():
            if name.endswith(".opf"):
                opf_name = name
                break

        if opf_name:
            opf = BeautifulSoup(z.read(opf_name), "lxml-xml")
            spine = opf.find_all("itemref")
            idrefs = [s.get("idref") for s in spine]
            id_map = {item.get("id"): item.get("href")
                      for item in opf.find_all("item")}
            reading_order = [id_map.get(ref) for ref in idrefs if id_map.get(ref)]
        else:
            reading_order = sorted([n for n in z.namelist()
                                    if n.endswith((".html", ".htm", ".xhtml"))])

        for fname in reading_order:
            if fname:
                try:
                    html = z.read(fname).decode("utf-8", errors="replace")
                    soup = BeautifulSoup(html, "lxml")
                    for tag in soup(["script", "style"]):
                        tag.decompose()
                    text = soup.get_text(separator="\n", strip=True)
                    if text.strip():
                        chapters.append(f"\n--- {fname} ---\n{text}")
                except Exception:
                    pass

    return "\n".join(chapters)


def load_rtf(path):
    """Load RTF — strips RTF control words."""
    from striprtf.striprtf import rtf_to_text
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return rtf_to_text(f.read())


def load_image(path):
    """OCR an image file (PNG, JPG, TIFF, BMP, etc.)."""
    import pytesseract
    from PIL import Image
    img = Image.open(path)
    text = pytesseract.image_to_string(img)
    return text.strip()


# ── Registry ────────────────────────────────────────────────────────────

LOADERS = {
    ".pdf":    load_pdf,
    ".txt":    load_txt,
    ".md":     load_txt,
    ".markdown": load_txt,
    ".docx":   load_docx,
    ".csv":    load_csv,
    ".json":   load_json,
    ".xlsx":   load_xlsx,
    ".xls":    load_xls,
    ".html":   load_html,
    ".htm":    load_html,
    ".xml":    load_xml,
    ".epub":   load_epub,
    ".rtf":    load_rtf,
    # Image formats (OCR)
    ".png":    load_image,
    ".jpg":    load_image,
    ".jpeg":   load_image,
    ".tiff":   load_image,
    ".tif":    load_image,
    ".bmp":    load_image,
    ".webp":   load_image,
}

SUPPORTED = ", ".join(sorted(LOADERS.keys()))


def load_document(path):
    """Load any supported document and return its text content."""
    ext = Path(path).suffix.lower()
    loader = LOADERS.get(ext)

    if loader is None:
        print(f"[!] Unsupported format '{ext}'. Trying as plain text.")
        return load_txt(path)

    print(f"[+] Loading {ext} file: {path}")
    return loader(path)
